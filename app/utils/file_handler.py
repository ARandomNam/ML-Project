"""
文件处理工具模块
支持PDF、DOCX等文件格式的文本提取
"""

import os
import io
import re
from typing import Optional, Union
import PyPDF2
from docx import Document


class FileHandler:
    """文件处理器类"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        从PDF文件中提取文本
        
        Args:
            file_content: PDF文件的字节内容
            
        Returns:
            提取的文本内容
        """
        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        从DOCX文件中提取文本
        
        Args:
            file_content: DOCX文件的字节内容
            
        Returns:
            提取的文本内容
        """
        try:
            docx_stream = io.BytesIO(file_content)
            doc = Document(docx_stream)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """
        从TXT文件中提取文本
        
        Args:
            file_content: TXT文件的字节内容
            
        Returns:
            提取的文本内容
        """
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'ascii', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用错误处理
            text = file_content.decode('utf-8', errors='ignore')
            return text.strip()
        
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ""
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """
        根据文件类型提取文本
        
        Args:
            file_content: 文件的字节内容
            filename: 文件名
            
        Returns:
            提取的文本内容
        """
        if not file_content:
            return ""
        
        # 获取文件扩展名
        _, ext = os.path.splitext(filename.lower())
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_content)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_content)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_content)
        else:
            # 尝试作为文本文件处理
            return self.extract_text_from_txt(file_content)
    
    def is_supported_format(self, filename: str) -> bool:
        """
        检查文件格式是否支持
        
        Args:
            filename: 文件名
            
        Returns:
            是否支持该格式
        """
        _, ext = os.path.splitext(filename.lower())
        return ext in self.supported_formats
    
    def validate_file_size(self, file_content: bytes, max_size_mb: int = 10) -> bool:
        """
        验证文件大小
        
        Args:
            file_content: 文件内容
            max_size_mb: 最大文件大小（MB）
            
        Returns:
            文件大小是否在允许范围内
        """
        if not file_content:
            return False
        
        file_size_mb = len(file_content) / (1024 * 1024)
        return file_size_mb <= max_size_mb
    
    def clean_extracted_text(self, text: str) -> str:
        """
        清理提取的文本
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        if not text:
            return ""
        
        # 去除多余的空行
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 去除行首尾空格
        lines = [line.strip() for line in text.split('\n')]
        
        # 去除空行
        lines = [line for line in lines if line]
        
        # 重新组合
        cleaned_text = '\n'.join(lines)
        
        return cleaned_text
    
    def process_uploaded_file(self, file_content: bytes, filename: str) -> dict:
        """
        处理上传的文件
        
        Args:
            file_content: 文件内容
            filename: 文件名
            
        Returns:
            处理结果字典
        """
        result = {
            'success': False,
            'text': '',
            'filename': filename,
            'error': None
        }
        
        try:
            # 验证文件格式
            if not self.is_supported_format(filename):
                result['error'] = f"不支持的文件格式。支持的格式：{', '.join(self.supported_formats)}"
                return result
            
            # 验证文件大小
            if not self.validate_file_size(file_content):
                result['error'] = "文件大小超过限制（最大10MB）"
                return result
            
            # 提取文本
            raw_text = self.extract_text_from_file(file_content, filename)
            
            if not raw_text:
                result['error'] = "无法从文件中提取文本内容"
                return result
            
            # 清理文本
            cleaned_text = self.clean_extracted_text(raw_text)
            
            result['success'] = True
            result['text'] = cleaned_text
            
        except Exception as e:
            result['error'] = f"处理文件时发生错误：{str(e)}"
        
        return result 